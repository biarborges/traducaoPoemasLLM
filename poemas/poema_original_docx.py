import csv
from docx import Document

def save_original_poem_to_word(csv_file, output_file):
    document = Document()  # Criando um novo documento Word
    
    with open(csv_file, mode="r", encoding="utf-8") as infile:
        reader = csv.DictReader(infile)  # Lendo o CSV
        
        for row in reader:
            # Adiciona o poema entre aspas
            poema_formatado = f'"{row["original_poem"]}"'
            document.add_paragraph(poema_formatado)  # Adiciona como parágrafo
            
            # Separador entre os poemas
            #document.add_paragraph("\n" + "-"*40 + "\n")
    
    document.save(output_file)  # Salvando o arquivo Word

# Exemplo de uso
csv_file = '../traducaoPoemasLLM/poemas/frances_ingles_poems.csv'  # Arquivo CSV original
output_file = '../traducaoPoemasLLM/poemas/frances_ingles_poems_apenas_frances.docx'  # Arquivo Word de saída

save_original_poem_to_word(csv_file, output_file)
print(f"Poemas salvos no formato desejado.")
