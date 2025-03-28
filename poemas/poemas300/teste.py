import csv
from docx import Document

def save_original_poem_to_word(csv_file, output_file):
    document = Document()  # Criando um novo documento Word
    
    with open(csv_file, mode="r", encoding="utf-8") as infile:
        reader = csv.DictReader(infile)  # Lendo o CSV
        
        for row in reader:
            document.add_paragraph(row["original_poem"])  # Adiciona cada poema como um parágrafo
            document.add_paragraph("\n" + "-"*40 + "\n")  # Separador entre os poemas
    
    document.save(output_file)  # Salvando o arquivo Word

# Exemplo de uso
csv_file = '../traducaoPoemasLLM/poemas/poemas300/frances_ingles_poems.csv'  # Arquivo CSV original
output_file = '../traducaoPoemasLLM/poemas/poemas300/frances_ingles_apenas_frances.docx'  # Arquivo Word de saída

save_original_poem_to_word(csv_file, output_file)
